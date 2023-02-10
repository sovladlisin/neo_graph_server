from django.shortcuts import render
from django.http import StreamingHttpResponse, HttpResponseRedirect, HttpResponse
from django.forms.models import model_to_dict
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
import datetime
from django.db.models import Q
from django.core.files.base import ContentFile
from .models import Resource
from .onthology_driver import Onthology
from.onthology_namespace import *

from core.settings import *

# API IMPORTS
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated, AllowAny


from docx import Document

@api_view(['POST', ])
@permission_classes((IsAuthenticated,))
def uploadFile(request):

 
    file_d = request.FILES['file']

    name = request.GET.get('name','')
    object_id = request.GET.get('object_id','')
    file_type = request.GET.get('file_type','')
    note = request.GET.get('note','')

    res_type = request.GET.get('res_type',None)

    o = Onthology(DB_URI,DB_USER, DB_PASSWORD)

    object_node = o.getEntityById(object_id)
    object_uri = object_node.get('uri')

    art_name = str(datetime.datetime.now().time())[:8]


    res = Resource()
    res.source.save(file_type + art_name + '.' + file_type,  ContentFile(file_d.read()))
    res.name = name
    res.resource_type = file_type
    res.save()


    r, visual = o.connectDigitalToResource(file_type, res.id,name,object_node.id, note,res_type )
    res.original_object_uri = r['uri']
    res.save()


    # collect file to return 

    response = {}
    response['resource'] = o.nodeToDict(visual)
    response['media'] = []
    response['genres'] = []
    response['lang'] = []
    response['events'] = []
    response['media_carrier'] = [
        {
            'file': {
                'name': name,
                'source': res.source.url,
                'id': res.id,
                'type': file_type
            }
        }
    ]

    o.close()



    return JsonResponse(response, safe=False)



@api_view(['DELETE', ])
@permission_classes((IsAuthenticated,))
def deleteFile(request):
    id = request.GET.get('id',None)
    if id is None:
        return HttpResponse(status=403)
    resource = Resource.objects.get(pk=id)
    resource.delete()
    return HttpResponse(status=200)

@api_view(['POST', ])
@permission_classes((IsAuthenticated,))
def cleanUp(request):
    res = 0
    if request.user.is_admin:

        o = Onthology(DB_URI,DB_USER, DB_PASSWORD)
        res = o.cleanUp()

        o.close()

    return Response({'counter': res})
    

@api_view(['POST', ])
@permission_classes((IsAuthenticated,))
def changeComments(request):

    data = json.loads(request.body.decode('utf-8'))
    comments = data.get('comments', None)
    commentary_uri = data.get('commentary_uri', None)

    commentary = ContentFile(b'')
    for comment in comments:
        temp = comment['text'] + '\n'
        commentary.write(temp.encode('utf-8'))



    resource_texts = Resource.objects.get(original_object_uri=commentary_uri)
    resource_texts.source.save(
        'commentary_' + str(resource_texts.pk) + '.txt', commentary)
    return HttpResponse(status=200)


@api_view(['POST', ])
@permission_classes((IsAuthenticated,))
def uploadDocxFirstTable(request):
    file_d = request.FILES['file']
    name = str(datetime.datetime.now().time())[:8]
    if file_d is None:
        return HttpResponse(status=403)
    docx_file = Document(file_d)
    tables = docx_file.tables
    info_table = tables[0]

    response = {}

    # название
    response['title'] = info_table.rows[0].cells[1].text if len(info_table.rows[0].cells[1].text) != 0 else ''

    # название на национальном языке
    response['lang_origin'] = info_table.rows[1].cells[1].text if len(info_table.rows[1].cells[1].text) != 0 else ''

    # язык
    response['lang'] = info_table.rows[2].cells[1].text if len(info_table.rows[2].cells[1].text) != 0 else ''

    # диалект
    response['dialect'] = info_table.rows[3].cells[1].text if len(info_table.rows[3].cells[1].text) != 0 else ''

    # говор
    response['speech'] = info_table.rows[4].cells[1].text if len(info_table.rows[4].cells[1].text) != 0 else ''

    # жанр
    response['genre'] = info_table.rows[5].cells[1].text if len(info_table.rows[5].cells[1].text) != 0 else ''

    # обряд
    response['obr'] = info_table.rows[6].cells[1].text if len(info_table.rows[6].cells[1].text) != 0 else ''

    # время записи
    response['time'] = info_table.rows[7].cells[1].text if len(info_table.rows[7].cells[1].text) != 0 else ''

    # место записи
    response['place'] = info_table.rows[8].cells[1].text if len(info_table.rows[8].cells[1].text) != 0 else ''

    # исполнитель
    response['permormed_by'] = info_table.rows[9].cells[1].text if len(info_table.rows[9].cells[1].text) != 0 else ''

    # собиратель
    response['collected_by'] = info_table.rows[10].cells[1].text if len(info_table.rows[10].cells[1].text) != 0 else ''

    # расшифровка аудиозаписи
    response['decrypted_by'] = info_table.rows[11].cells[1].text if len(info_table.rows[11].cells[1].text) != 0 else ''

    # нотирование
    response['notation_by'] = info_table.rows[12].cells[1].text if len(info_table.rows[12].cells[1].text) != 0 else ''

    # перевод на русский язык
    response['transalted_by'] = info_table.rows[13].cells[1].text if len(info_table.rows[13].cells[1].text) != 0 else ''

    # редактор перевода
    response['editor'] = info_table.rows[14].cells[1].text if len(info_table.rows[14].cells[1].text) != 0 else ''

    # редактор национального текста
    response['redactor'] = info_table.rows[15].cells[1].text if len(info_table.rows[15].cells[1].text) != 0 else ''

    # подготовка клмментариев
    response['commantator'] = info_table.rows[16].cells[1].text if len(info_table.rows[16].cells[1].text) != 0 else ''

    # опубликовано
    response['published'] = info_table.rows[17].cells[1].text if len(info_table.rows[17].cells[1].text) != 0 else ''

    # место хранения
    response['place_storage'] = info_table.rows[18].cells[1].text if len(info_table.rows[18].cells[1].text) != 0 else ''

    # варианты
    response['variants'] = info_table.rows[19].cells[1].text if len(info_table.rows[19].cells[1].text) != 0 else ''

    # дополнительная иформация
    response['note'] = info_table.rows[20].cells[1].text if len(info_table.rows[20].cells[1].text) != 0 else ''

    return Response(response)
    main_table = tables[1]

    temp = ''
    original = ContentFile(b'')
    translation = ContentFile(b'')
    commentary = ContentFile(b'')

    for cell in main_table.columns[0].cells:
        temp = cell.text + '\n'
        original.write(temp.encode('utf-8'))

    for cell in main_table.columns[1].cells:
        temp = cell.text + '\n'
        translation.write(temp.encode('utf-8'))

    for cell in main_table.columns[2].cells:
        temp = cell.text + '\n'
        commentary.write(temp.encode('utf-8'))


    original_r = Resource()
    original_r.source.save('original_' + name,  original)

    trans_r = Resource()
    trans_r.source.save('translation_' + name,  translation)

    comment_r = Resource()
    comment_r.source.save('comment_' + name,  commentary)




    original_r.save()
    trans_r.save()
    comment_r.save()


    return HttpResponse(status=200)

@api_view(['POST', ])
@permission_classes((IsAuthenticated,))
def uploadDocx(request):
    file_d = request.FILES['file']
    data = request.data['data']
    data = json.loads(data)
    
    
    node = data['node']
    corpus_id = data['corpus_id']




    name = str(datetime.datetime.now().time())[:8]
    if file_d is None:
        return HttpResponse(status=403)
    docx_file = Document(file_d)
    tables = docx_file.tables
    main_table = tables[1]

    temp = ''
    original = ContentFile(b'')
    translation = ContentFile(b'')
    commentary = ContentFile(b'')

    for cell in main_table.columns[0].cells:
        temp = cell.text + '\n'
        original.write(temp.encode('utf-8'))

    for cell in main_table.columns[1].cells:
        temp = cell.text + '\n'
        translation.write(temp.encode('utf-8'))

    for cell in main_table.columns[2].cells:
        temp = cell.text + '\n'
        commentary.write(temp.encode('utf-8'))

    original_r = Resource()
    original_r.source.save('original_' + name,  original)

    trans_r = Resource()
    trans_r.source.save('translation_' + name,  translation)

    comment_r = Resource()
    comment_r.source.save('comment_' + name,  commentary)

    o = Onthology(DB_URI,DB_USER, DB_PASSWORD)
    origin_node_uri, transaltion_node_uri, commentary_node_uri, origin_node = o.createText(node, corpus_id, original_r.pk, trans_r.pk, comment_r.pk)

    original_r.name = 'original_' + name
    original_r.original_object_uri = origin_node_uri
    original_r.resource_type = 'text'

    trans_r.name = 'translation_' + name
    trans_r.original_object_uri = transaltion_node_uri
    trans_r.resource_type = 'text'

    comment_r.name = 'comment_' + name
    comment_r.original_object_uri = commentary_node_uri
    comment_r.resource_type = 'text'


    original_r.save()
    trans_r.save()
    comment_r.save()


    # return response to update 

    response = {}
    response['resource'] = origin_node
    response['media'] = []
    response['genres'] = []
    response['lang'] = []
    response['events'] = []
    response['media_carrier'] = [
        {
            'file': {
                'name': 'original_' + name,
                'source': original_r.source.url,
                'id':  original_r.id,
                'type': 'text'
            }
        }
    ]

    o.close()
    
    # print('IDIDIDIDIDI:', created_node.id)
    return Response(response)